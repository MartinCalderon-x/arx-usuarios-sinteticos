"""Document processor service for extracting text from multiple file formats."""
import io
import csv
from typing import Optional


class DocumentProcessor:
    """Service for extracting text from documents (PDF, DOCX, TXT, CSV)."""

    # Supported MIME types
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/csv": "csv",
        "application/csv": "csv",
    }

    # File size limits
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB per file
    MAX_TOTAL_SIZE = 50 * 1024 * 1024  # 50MB total
    MAX_FILES = 5

    def __init__(self):
        self._pdf_available = False
        self._docx_available = False
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which document processing libraries are available."""
        try:
            import pypdf
            self._pdf_available = True
        except ImportError:
            pass

        try:
            import docx
            self._docx_available = True
        except ImportError:
            pass

    def is_supported(self, content_type: str) -> bool:
        """Check if a content type is supported."""
        return content_type in self.SUPPORTED_TYPES

    def get_file_type(self, content_type: str) -> Optional[str]:
        """Get file type from content type."""
        return self.SUPPORTED_TYPES.get(content_type)

    async def extract_text(self, file_data: bytes, content_type: str, filename: str = "") -> dict:
        """
        Extract text from a file.

        Args:
            file_data: Raw file bytes
            content_type: MIME type of the file
            filename: Original filename (for fallback type detection)

        Returns:
            dict with:
                - text: Extracted text content
                - success: Whether extraction was successful
                - error: Error message if failed
                - char_count: Number of characters extracted
        """
        file_type = self.get_file_type(content_type)

        # Fallback to extension-based detection
        if not file_type and filename:
            ext = filename.lower().split(".")[-1] if "." in filename else ""
            type_map = {"pdf": "pdf", "docx": "docx", "txt": "txt", "csv": "csv"}
            file_type = type_map.get(ext)

        if not file_type:
            return {
                "text": "",
                "success": False,
                "error": f"Tipo de archivo no soportado: {content_type}",
                "char_count": 0
            }

        try:
            if file_type == "pdf":
                text = await self._extract_pdf(file_data)
            elif file_type == "docx":
                text = await self._extract_docx(file_data)
            elif file_type == "txt":
                text = await self._extract_txt(file_data)
            elif file_type == "csv":
                text = await self._extract_csv(file_data)
            else:
                return {
                    "text": "",
                    "success": False,
                    "error": f"Tipo de archivo no implementado: {file_type}",
                    "char_count": 0
                }

            return {
                "text": text,
                "success": True,
                "error": None,
                "char_count": len(text)
            }

        except Exception as e:
            return {
                "text": "",
                "success": False,
                "error": f"Error al procesar archivo: {str(e)}",
                "char_count": 0
            }

    async def _extract_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF file."""
        if not self._pdf_available:
            raise ImportError("pypdf no está instalado. Instalar con: pip install pypdf")

        import pypdf

        pdf_file = io.BytesIO(file_data)
        reader = pypdf.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Página {page_num + 1} ---\n{page_text}")

        return "\n\n".join(text_parts)

    async def _extract_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX file."""
        if not self._docx_available:
            raise ImportError("python-docx no está instalado. Instalar con: pip install python-docx")

        import docx

        docx_file = io.BytesIO(file_data)
        doc = docx.Document(docx_file)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    async def _extract_txt(self, file_data: bytes) -> str:
        """Extract text from plain text file."""
        # Try different encodings
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                return file_data.decode(encoding)
            except UnicodeDecodeError:
                continue

        # Last resort: decode with errors ignored
        return file_data.decode("utf-8", errors="ignore")

    async def _extract_csv(self, file_data: bytes) -> str:
        """Extract text from CSV file, formatted for readability."""
        text = await self._extract_txt(file_data)

        # Parse CSV and format as readable text
        csv_file = io.StringIO(text)
        reader = csv.reader(csv_file)

        rows = list(reader)
        if not rows:
            return ""

        # Format as readable table
        text_parts = []
        headers = rows[0] if rows else []

        for i, row in enumerate(rows):
            if i == 0 and headers:
                text_parts.append("ENCABEZADOS: " + " | ".join(headers))
                text_parts.append("-" * 50)
            else:
                # Format as key-value pairs for better LLM understanding
                if headers and len(row) == len(headers):
                    formatted = ", ".join(f"{headers[j]}: {row[j]}" for j in range(len(row)))
                    text_parts.append(f"Registro {i}: {formatted}")
                else:
                    text_parts.append(f"Registro {i}: {' | '.join(row)}")

        return "\n".join(text_parts)

    async def process_multiple_files(
        self,
        files: list[tuple[bytes, str, str]]
    ) -> dict:
        """
        Process multiple files and combine their text.

        Args:
            files: List of tuples (file_data, content_type, filename)

        Returns:
            dict with:
                - combined_text: All extracted text combined
                - files_processed: Number of successfully processed files
                - files_failed: Number of failed files
                - errors: List of error messages
                - total_chars: Total character count
        """
        if len(files) > self.MAX_FILES:
            return {
                "combined_text": "",
                "files_processed": 0,
                "files_failed": len(files),
                "errors": [f"Máximo {self.MAX_FILES} archivos permitidos"],
                "total_chars": 0
            }

        # Check total size
        total_size = sum(len(f[0]) for f in files)
        if total_size > self.MAX_TOTAL_SIZE:
            return {
                "combined_text": "",
                "files_processed": 0,
                "files_failed": len(files),
                "errors": [f"Tamaño total excede {self.MAX_TOTAL_SIZE // (1024*1024)}MB"],
                "total_chars": 0
            }

        text_parts = []
        errors = []
        files_processed = 0

        for i, (file_data, content_type, filename) in enumerate(files):
            # Check individual file size
            if len(file_data) > self.MAX_FILE_SIZE:
                errors.append(f"{filename}: Excede {self.MAX_FILE_SIZE // (1024*1024)}MB")
                continue

            result = await self.extract_text(file_data, content_type, filename)

            if result["success"]:
                text_parts.append(f"=== ARCHIVO: {filename} ===\n{result['text']}")
                files_processed += 1
            else:
                errors.append(f"{filename}: {result['error']}")

        combined_text = "\n\n" + "="*60 + "\n\n".join(text_parts) if text_parts else ""

        return {
            "combined_text": combined_text,
            "files_processed": files_processed,
            "files_failed": len(files) - files_processed,
            "errors": errors,
            "total_chars": len(combined_text)
        }


# Singleton instance
document_processor = DocumentProcessor()
